{\rtf1\ansi\ansicpg1252\cocoartf2865
\cocoatextscaling0\cocoaplatform0{\fonttbl\f0\fmodern\fcharset0 Courier;}
{\colortbl;\red255\green255\blue255;\red0\green0\blue0;}
{\*\expandedcolortbl;;\cssrgb\c0\c0\c0;}
\paperw11900\paperh16840\margl1440\margr1440\vieww11520\viewh8400\viewkind0
\deftab720
\pard\pardeftab720\partightenfactor0

\f0\fs26 \cf0 \expnd0\expndtw0\kerning0
\outl0\strokewidth0 \strokec2 import streamlit as st\
import sqlite3\
import pandas as pd\
from datetime import date, datetime\
from docxtpl import DocxTemplate\
from docx import Document\
from docx.shared import Pt\
from docx.enum.text import WD_ALIGN_PARAGRAPH\
import os\
import io\
\
# \'e2\'94\'80\'e2\'94\'80 CONFIG \'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\
DB_PATH = "clients.db"\
TEMPLATE_PATH = "execution_petition_TEMPLATE.docx"\
\
st.set_page_config(\
    page_title="Execution Petition Manager",\
    page_icon="\'e2\'9a\'96\'ef\'b8\uc0\u143 ",\
    layout="wide"\
)\
\
# \'e2\'94\'80\'e2\'94\'80 DATABASE HELPERS \'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\
def get_connection():\
    return sqlite3.connect(DB_PATH)\
\
def init_db():\
    conn = get_connection()\
    conn.execute('''CREATE TABLE IF NOT EXISTS clients (\
        id INTEGER PRIMARY KEY AUTOINCREMENT,\
        case_number TEXT UNIQUE,\
        petitioner_name TEXT,\
        petitioner_address TEXT,\
        respondent_name TEXT,\
        respondent_address TEXT,\
        marriage_date TEXT,\
        decree_date TEXT,\
        children_info TEXT,\
        decree_amount TEXT,\
        court_costs TEXT,\
        execution_mode TEXT,\
        filing_deadline TEXT,\
        prepared_by TEXT,\
        status TEXT DEFAULT "Active"\
    )''')\
    conn.commit()\
    conn.close()\
\
def load_clients(search=""):\
    conn = get_connection()\
    if search:\
        query = """SELECT * FROM clients WHERE\
            petitioner_name LIKE ? OR respondent_name LIKE ? OR\
            case_number LIKE ? OR status LIKE ?"""\
        s = f"%\{search\}%"\
        df = pd.read_sql_query(query, conn, params=(s, s, s, s))\
    else:\
        df = pd.read_sql_query("SELECT * FROM clients ORDER BY filing_deadline ASC", conn)\
    conn.close()\
    return df\
\
def add_client(data: dict):\
    conn = get_connection()\
    conn.execute('''INSERT INTO clients\
        (case_number, petitioner_name, petitioner_address, respondent_name,\
         respondent_address, marriage_date, decree_date, children_info,\
         decree_amount, court_costs, execution_mode, filing_deadline,\
         prepared_by, status)\
        VALUES (:case_number,:petitioner_name,:petitioner_address,:respondent_name,\
                :respondent_address,:marriage_date,:decree_date,:children_info,\
                :decree_amount,:court_costs,:execution_mode,:filing_deadline,\
                :prepared_by,:status)''', data)\
    conn.commit()\
    conn.close()\
\
def delete_client(case_id: int):\
    conn = get_connection()\
    conn.execute("DELETE FROM clients WHERE id=?", (case_id,))\
    conn.commit()\
    conn.close()\
\
def ensure_template():\
    """Create the Word template if it doesn't exist."""\
    if os.path.exists(TEMPLATE_PATH):\
        return\
    doc = Document()\
    doc.styles["Normal"].font.name = "Times New Roman"\
    doc.styles["Normal"].font.size = Pt(11)\
    t = doc.add_heading("EXECUTION PETITION", 0)\
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER\
    s = doc.add_paragraph("IN THE FAMILY COURT")\
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER\
    s2 = doc.add_paragraph("Case No: \{\{case_number\}\}")\
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER\
    doc.add_paragraph("")\
    doc.add_heading("PARTIES", level=1)\
    p = doc.add_paragraph()\
    p.add_run("Decree Holder (Petitioner): ").bold = True\
    p.add_run("\{\{petitioner_name\}\}, residing at \{\{petitioner_address\}\}")\
    p2 = doc.add_paragraph()\
    p2.add_run("Judgment Debtor (Respondent): ").bold = True\
    p2.add_run("\{\{respondent_name\}\}, residing at \{\{respondent_address\}\}")\
    doc.add_paragraph("")\
    doc.add_heading("PARTICULARS OF EXECUTION", level=1)\
    table = doc.add_table(rows=10, cols=2)\
    table.style = "Table Grid"\
    rows_data = [\
        ("1.  Case Number", "\{\{case_number\}\}"),\
        ("2.  Petitioner vs Respondent", "\{\{petitioner_name\}\} vs \{\{respondent_name\}\}"),\
        ("3.  Date of Marriage", "\{\{marriage_date\}\}"),\
        ("4.  Date of Divorce Decree", "\{\{decree_date\}\}"),\
        ("5.  Children", "\{\{children_info\}\}"),\
        ("6.  Relief / Amount Granted", "\{\{decree_amount\}\}"),\
        ("7.  Court Costs Allowed", "\{\{court_costs\}\}"),\
        ("8.  Execution Against", "\{\{respondent_name\}\}"),\
        ("9.  Mode of Execution", "\{\{execution_mode\}\}"),\
        ("10. Filing Deadline", "\{\{filing_deadline\}\}"),\
    ]\
    for i, (label, value) in enumerate(rows_data):\
        row = table.rows[i]\
        row.cells[0].text = label\
        row.cells[1].text = value\
        row.cells[0].paragraphs[0].runs[0].bold = True\
    doc.add_paragraph("")\
    doc.add_heading("PRAYER", level=1)\
    doc.add_paragraph(\
        "The Decree Holder respectfully requests that this Hon'ble Court execute the "\
        "order dated \{\{decree_date\}\} against \{\{respondent_name\}\} and provide all necessary assistance."\
    )\
    doc.add_paragraph("")\
    doc.add_paragraph("Date: \{\{today_date\}\}")\
    doc.add_paragraph("Prepared by: \{\{prepared_by\}\}")\
    doc.add_paragraph("")\
    doc.add_paragraph("_" * 40)\
    doc.add_paragraph("Signature of Decree Holder / Authorized Representative")\
    doc.save(TEMPLATE_PATH)\
\
def generate_doc(case: dict) -> bytes:\
    ensure_template()\
    tpl = DocxTemplate(TEMPLATE_PATH)\
    context = \{k: (v if v else "-") for k, v in case.items()\}\
    context["today_date"] = str(date.today())\
    tpl.render(context)\
    buf = io.BytesIO()\
    tpl.save(buf)\
    buf.seek(0)\
    return buf.read()\
\
# \'e2\'94\'80\'e2\'94\'80 INIT \'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\
init_db()\
ensure_template()\
\
# \'e2\'94\'80\'e2\'94\'80 SIDEBAR \'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\'e2\'94\'80\
st.sidebar.image("https://img.icons8.com/color/96/scales--v1.png", width=60)\
st.sidebar.title("\'e2\'9a\'96\'ef\'b8\uc0\u143  Petition Manager")\
page = st.sidebar.radio("Navigate", ["\'f0\'9f\'93\'8b Client Table", "\'e2\'9e\'95 Add New Case", "\'f0\'9f\'93\'84 Generate Document"])\
st.sidebar.markdown("---")\
st.sidebar.caption("Built for legal document automation")\
\
# \'e2\'95\uc0\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \
# PAGE 1 \'e2\'80\'94 CLIENT TABLE\
# \'e2\'95\uc0\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \
if page == "\'f0\'9f\'93\'8b Client Table":\
    st.title("\'f0\'9f\'93\'8b Client Case Table")\
\
    col1, col2 = st.columns([3, 1])\
    with col1:\
        search = st.text_input("\'f0\'9f\'94\uc0\u141  Search by name, case number, or status", "")\
    with col2:\
        st.markdown("<br>", unsafe_allow_html=True)\
        refresh = st.button("\'f0\'9f\'94\'84 Refresh")\
\
    df = load_clients(search)\
\
    if df.empty:\
        st.info("No clients found. Add a new case from the sidebar.")\
    else:\
        # Deadline urgency highlighting\
        today = date.today()\
        def urgency(deadline_str):\
            try:\
                d = datetime.strptime(deadline_str, "%Y-%m-%d").date()\
                diff = (d - today).days\
                if diff < 0:   return "\'f0\'9f\'94\'b4 OVERDUE"\
                elif diff <= 2: return "\'f0\'9f\'94\'b4 Urgent"\
                elif diff <= 7: return "\'f0\'9f\'9f\'a0 This Week"\
                else:           return "\'f0\'9f\'9f\'a2 On Track"\
            except:\
                return "\'e2\'80\'94"\
\
        df["\'e2\uc0\u143 \'b0 Urgency"] = df["filing_deadline"].apply(urgency)\
        display_cols = ["case_number", "petitioner_name", "respondent_name",\
                        "filing_deadline", "\'e2\uc0\u143 \'b0 Urgency", "status", "execution_mode"]\
        st.dataframe(df[display_cols], use_container_width=True, hide_index=True)\
\
        st.markdown(f"**Total cases: \{len(df)\}**")\
\
        # Stats\
        c1, c2, c3 = st.columns(3)\
        overdue = df[df["\'e2\uc0\u143 \'b0 Urgency"].str.contains("OVERDUE|Urgent")].shape[0]\
        this_week = df[df["\'e2\uc0\u143 \'b0 Urgency"] == "\'f0\'9f\'9f\'a0 This Week"].shape[0]\
        active = df[df["status"] == "Active"].shape[0]\
        c1.metric("\'f0\'9f\'94\'b4 Urgent / Overdue", overdue)\
        c2.metric("\'f0\'9f\'9f\'a0 Due This Week", this_week)\
        c3.metric("\'e2\'9c\'85 Active Cases", active)\
\
        # Delete\
        st.markdown("---")\
        with st.expander("\'f0\'9f\'97\'91\'ef\'b8\uc0\u143  Delete a Case"):\
            case_ids = df[["id", "case_number", "petitioner_name"]].copy()\
            case_ids["label"] = case_ids["case_number"] + " \'e2\'80\'93 " + case_ids["petitioner_name"]\
            chosen = st.selectbox("Select case to delete", case_ids["label"].tolist())\
            if st.button("Delete Selected Case", type="secondary"):\
                cid = case_ids[case_ids["label"] == chosen]["id"].values[0]\
                delete_client(int(cid))\
                st.success(f"Deleted: \{chosen\}")\
                st.rerun()\
\
# \'e2\'95\uc0\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \
# PAGE 2 \'e2\'80\'94 ADD NEW CASE\
# \'e2\'95\uc0\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \
elif page == "\'e2\'9e\'95 Add New Case":\
    st.title("\'e2\'9e\'95 Add New Client Case")\
\
    with st.form("new_case_form", clear_on_submit=True):\
        st.subheader("\'f0\'9f\'93\uc0\u129  Case Info")\
        c1, c2 = st.columns(2)\
        case_number  = c1.text_input("Case Number *", placeholder="FAM-2026-004")\
        prepared_by  = c2.text_input("Prepared By *", placeholder="Your name")\
        status       = c1.selectbox("Status", ["Active", "Completed", "On Hold"])\
        filing_deadline = c2.date_input("Filing Deadline *", min_value=date.today())\
\
        st.subheader("\'f0\'9f\'91\'a9 Petitioner (Decree Holder)")\
        c3, c4 = st.columns(2)\
        petitioner_name    = c3.text_input("Full Name *")\
        petitioner_address = c4.text_input("Address *")\
\
        st.subheader("\'f0\'9f\'91\'a8 Respondent (Judgment Debtor)")\
        c5, c6 = st.columns(2)\
        respondent_name    = c5.text_input("Full Name *", key="resp_name")\
        respondent_address = c6.text_input("Address *", key="resp_addr")\
\
        st.subheader("\'f0\'9f\'93\'85 Marriage & Decree")\
        c7, c8 = st.columns(2)\
        marriage_date = c7.date_input("Date of Marriage")\
        decree_date   = c8.date_input("Date of Divorce Decree")\
\
        st.subheader("\'e2\'9a\'96\'ef\'b8\uc0\u143  Execution Details")\
        children_info  = st.text_input("Children Info", placeholder="2 children: Ali (7), Sara (5)")\
        c9, c10 = st.columns(2)\
        decree_amount  = c9.text_input("Relief / Amount Granted", placeholder="5,000,000 UZS/month")\
        court_costs    = c10.text_input("Court Costs Allowed", placeholder="250,000 UZS")\
        execution_mode = st.selectbox("Mode of Execution", [\
            "Salary Attachment", "Bank Account Garnishment",\
            "Property Attachment", "Other"\
        ])\
\
        submitted = st.form_submit_button("\'f0\'9f\'92\'be Save Case", type="primary")\
        if submitted:\
            if not all([case_number, petitioner_name, respondent_name, prepared_by]):\
                st.error("Please fill in all required (*) fields.")\
            else:\
                try:\
                    add_client(\{\
                        "case_number": case_number,\
                        "petitioner_name": petitioner_name,\
                        "petitioner_address": petitioner_address,\
                        "respondent_name": respondent_name,\
                        "respondent_address": respondent_address,\
                        "marriage_date": str(marriage_date),\
                        "decree_date": str(decree_date),\
                        "children_info": children_info,\
                        "decree_amount": decree_amount,\
                        "court_costs": court_costs,\
                        "execution_mode": execution_mode,\
                        "filing_deadline": str(filing_deadline),\
                        "prepared_by": prepared_by,\
                        "status": status,\
                    \})\
                    st.success(f"\'e2\'9c\'85 Case \{case_number\} saved successfully!")\
                    st.balloons()\
                except Exception as e:\
                    if "UNIQUE" in str(e):\
                        st.error(f"Case number \{case_number\} already exists.")\
                    else:\
                        st.error(f"Error: \{e\}")\
\
# \'e2\'95\uc0\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \
# PAGE 3 \'e2\'80\'94 GENERATE DOCUMENT\
# \'e2\'95\uc0\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \'e2\'95\u144 \
elif page == "\'f0\'9f\'93\'84 Generate Document":\
    st.title("\'f0\'9f\'93\'84 Generate Execution Petition")\
\
    df = load_clients()\
    if df.empty:\
        st.warning("No cases in the database. Please add a case first.")\
    else:\
        df["label"] = df["case_number"] + " \'e2\'80\'94 " + df["petitioner_name"] + " vs " + df["respondent_name"]\
        selected_label = st.selectbox("Select a case to generate document for:", df["label"].tolist())\
        selected_row = df[df["label"] == selected_label].iloc[0].to_dict()\
\
        st.markdown("### \'f0\'9f\'93\uc0\u157  Case Preview")\
        preview_cols = \{\
            "Case Number": selected_row["case_number"],\
            "Petitioner": selected_row["petitioner_name"],\
            "Respondent": selected_row["respondent_name"],\
            "Decree Date": selected_row["decree_date"],\
            "Filing Deadline": selected_row["filing_deadline"],\
            "Mode of Execution": selected_row["execution_mode"],\
            "Status": selected_row["status"],\
        \}\
        for k, v in preview_cols.items():\
            c1, c2 = st.columns([1, 2])\
            c1.markdown(f"**\{k\}**")\
            c2.markdown(v or "\'e2\'80\'94")\
\
        st.markdown("---")\
        if st.button("\'e2\'ac\'87\'ef\'b8\uc0\u143  Generate & Download Document", type="primary"):\
            with st.spinner("Generating document..."):\
                doc_bytes = generate_doc(selected_row)\
            filename = f"petition_\{selected_row['case_number'].replace('-', '_')\}.docx"\
            st.download_button(\
                label="\'f0\'9f\'93\'a5 Click here to download",\
                data=doc_bytes,\
                file_name=filename,\
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"\
            )\
            st.success(f"\'e2\'9c\'85 Document ready: \{filename\}")\
}